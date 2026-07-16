"""Generate per-theme Word reference documents from the epyson themes.

Pandoc's ``docx`` writer copies the named paragraph/character styles, the
default font and the page geometry from a ``--reference-doc``; it ignores the
document *body*, so the reference only needs correct **styles**.  This script
builds one clean reference per bundled theme straight from a fresh
``docx.Document()`` — python-docx's default template already carries every
style Pandoc maps onto (Normal, Heading 1-9, Title, Subtitle, Caption, Quote,
…) and ships with **no** header, footer or ``word/media`` image.

For every theme it re-skins those styles from the :class:`Theme`'s
``css_vars`` so the exported Word document carries the same identity as the
live HTML preview (mirroring what ``make_reference_pptx`` does for the deck
colour scheme and fonts), then writes a clean reference with NO ANM logo, NO
header/footer image and NO ``ANM-FG`` document code.

Run from the project root::

    python src/epy_reports/_core/_packaging/make_reference_docx.py

It writes ``src/epy_reports/_config/_assets/reference_docx/<theme>.docx``
for every bundled theme (the build-time dependency ``python-docx`` is not
required at runtime — Pandoc consumes the generated files).
"""

from __future__ import annotations

import sys
from pathlib import Path

# Repo root: four levels above this file (_packaging -> _core ->
# epy_reports -> src -> root).
ROOT = Path(__file__).resolve().parents[4]
sys.path.insert(0, str(ROOT / "src"))

from docx import Document  # noqa: E402
from docx.shared import Inches, Pt, RGBColor  # noqa: E402

from epy_reports import themes  # noqa: E402
from epy_reports.themes_base import Theme  # noqa: E402

OUT_DIR = (
    ROOT / "src" / "epy_reports" / "_config" / "_assets" / "reference_docx"
)

# US Letter and conventional 1 inch margins for the reference page geometry.
_LETTER_W = Inches(8.5)
_LETTER_H = Inches(11)
_MARGIN = Inches(1)

# Heading style name -> (size css var, weight css var). Pandoc maps Markdown
# ``#``..``######`` onto Word "Heading 1".."Heading 6"; "Heading 7"-"9" are
# left at python-docx defaults (Pandoc never emits past level 6).
_HEADINGS = {
    "Heading 1": ("h1-size", "h1-weight"),
    "Heading 2": ("h2-size", "h2-weight"),
    "Heading 3": ("h3-size", "h3-weight"),
    "Heading 4": ("h4-size", "h4-weight"),
    "Heading 5": ("h5-size", "h5-weight"),
    "Heading 6": ("h6-size", "h6-weight"),
}


def _first_font(value: str, default: str) -> str:
    """Return the first family name from a CSS ``font-family`` list."""
    if not value:
        return default
    first = value.split(",")[0].strip().strip('"').strip("'")
    return first or default


def _pt(value: str, default: float) -> float:
    """Parse a ``"12pt"`` (or ``"12"``) css size into a float point value."""
    text = (value or "").strip().lower().removesuffix("pt").strip()
    try:
        return float(text)
    except ValueError:
        return default


def _bold(weight: str) -> bool:
    """Return True when a css font-weight reads as bold (>= 600 or keyword)."""
    text = (weight or "").strip().lower()
    if text in {"bold", "bolder"}:
        return True
    try:
        return int(text) >= 600
    except ValueError:
        return False


def _rgb(value: str, default: str = "000000") -> RGBColor:
    """Build an :class:`RGBColor` from a ``#RRGGBB`` css colour value."""
    text = (value or "").strip().lstrip("#")
    if len(text) != 6:
        text = default
    try:
        return RGBColor.from_string(text.upper())
    except ValueError:
        return RGBColor.from_string(default)


def build_reference(theme: Theme, target: Path) -> None:
    """Write a clean, theme-styled reference ``.docx`` for ``theme``.

    Starts from a fresh ``Document()`` (default styles, no header/footer, no
    media) and re-skins the named styles Pandoc copies, so the reference
    reflects the theme without carrying any ANM branding.
    """
    css = theme.css_vars
    doc = Document()

    # ---- page geometry: US Letter + 1in margins --------------------
    for section in doc.sections:
        section.page_width = _LETTER_W
        section.page_height = _LETTER_H
        section.top_margin = _MARGIN
        section.bottom_margin = _MARGIN
        section.left_margin = _MARGIN
        section.right_margin = _MARGIN

    body_font = _first_font(css.get("font-family-text", ""), "Calibri")
    head_font = _first_font(css.get("font-family-headings", ""), body_font)
    body_pt = _pt(css.get("body-size", ""), 11.0)
    fg = _rgb(css.get("fg", ""), "000000")
    head_color = _rgb(css.get("heading-color", css.get("fg", "")), "000000")

    # ---- Normal: body font / size / text colour -------------------
    normal = doc.styles["Normal"]
    normal.font.name = body_font
    normal.font.size = Pt(body_pt)
    normal.font.color.rgb = fg

    # ---- Body Text / No Spacing inherit the body identity ---------
    for name in ("Body Text", "No Spacing"):
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        style.font.name = body_font
        style.font.color.rgb = fg

    # ---- Heading 1..6: heading font / size / weight / colour ------
    for name, (size_var, weight_var) in _HEADINGS.items():
        style = doc.styles[name]
        style.font.name = head_font
        style.font.size = Pt(_pt(css.get(size_var, ""), body_pt + 2))
        style.font.bold = _bold(css.get(weight_var, "700"))
        style.font.color.rgb = head_color

    # ---- Title / Subtitle: heading family, scaled from H1 ---------
    h1_pt = _pt(css.get("h1-size", ""), body_pt + 12)
    for name, scale, bold in (("Title", 1.6, True), ("Subtitle", 1.15, False)):
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        style.font.name = head_font
        style.font.size = Pt(round(h1_pt * scale, 1))
        style.font.bold = bold
        style.font.color.rgb = head_color

    # ---- Caption / Quote follow the muted + accent css vars -------
    try:
        caption = doc.styles["Caption"]
        caption.font.name = body_font
        caption.font.size = Pt(_pt(css.get("caption-size", ""), body_pt - 1))
        caption.font.color.rgb = _rgb(css.get("fg-muted", ""), "606060")
    except KeyError:
        pass
    for name in ("Quote", "Intense Quote"):
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        style.font.name = body_font
        accent = css.get("quote-rule", css.get("link", ""))
        style.font.color.rgb = _rgb(accent, "000000")

    doc.save(str(target))


def main() -> int:
    """Generate a clean reference document for every bundled theme."""
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for theme_id, theme in themes.THEMES.items():
        target = OUT_DIR / f"{theme_id}.docx"
        build_reference(theme, target)
        print(f"wrote {target.relative_to(ROOT)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
